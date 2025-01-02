from docx import Document
from datetime import datetime
from dateutil.relativedelta import relativedelta
import os
import pathlib
import json
import sys
import shutil

working_directory = pathlib.Path(os.getcwd())
project_path = pathlib.Path(os.path.dirname(os.path.abspath(__file__)))

cvc_code_json = project_path / "CVC Codes.json"
insurance_emails_json = project_path / "Insurance Emails.json"

# Templates
# Regular Template to send to OPINS
template_word_path_opins = project_path / "Template - SHORT - Individual.docx"

# Templates to send to CINS - UM - Uninsured Motorist. && UM AUTH
template_word_path_cins_a = project_path / "Template UM - SHORT.docx"
template_word_path_cins_b = project_path / "Template  Formal Demand for UMA.docx"

file_template_data = project_path / "template info - SHORT - Individual.txt"
file_template_source = working_directory / (working_directory.name + ".txt")

# Function to ensure the target file exists
def ensure_file_exists(src, dst):
    # Check if the target file exists
    if not os.path.exists(dst):
        print(f"Target file '{dst}' does not exist. Copying from source.")
        
        # Check if the source file exists
        if not os.path.exists(src):
            print(f"Source file '{src}' does not exist. Aborting operation.")
            return
        
        # Copy the source file to the destination
        shutil.copy(src, dst)
        print(f"File copied successfully from '{src}' to '{dst}'.")
    

def check_and_warn_if_file_exists(file_path):
    if file_path.exists():
        return False
    return True


def read_json_file(json_path):
    with open(json_path, "r", encoding="utf-8") as file:
        return json.load(file)  # Return the loaded JSON data

def update_json_file(json_path, cvc_json, cvc_code):
    # Prompt the user for the text corresponding to the CVC code
    if cvc_code:
        cvc_json[cvc_code] = input(f"Please enter the text for the CVC code '{cvc_code.upper()}' missing in the database: ")

        # Save the updated dictionary back to the JSON file
        with open(json_path, "w", encoding="utf-8") as file:
            json.dump(cvc_json, file, ensure_ascii=False, indent=4)

        # Return the updated dictionary
        return cvc_json
    print("CVC Code field is empty. Fill it before continue, exiting...")
    sys.exit(1)

def create_string(cvc_codes):
    cvc_text = ""
    cvc_json = read_json_file(cvc_code_json)

    if "," in cvc_codes:
        for index, cvc_code in enumerate(cvc_codes.split(", ")):
            while True:
                try:
                    cvc_text_code = cvc_json[cvc_code]
                    break
                except KeyError:
                    cvc_json = update_json_file(cvc_code_json, cvc_json, cvc_code)

            cvc_text += 'California Vehicle Code ' + cvc_code + ' "' + cvc_text_code + '"'
            if index != len(cvc_codes.split(", "))-1:
                cvc_text += ", "
            else:
                cvc_text += "."
    else:
        while True:
            try:
                cvc_text_code = cvc_json[cvc_codes]
                break
            except KeyError:
                cvc_json = update_json_file(cvc_code_json, cvc_json, cvc_codes)
        cvc_text = 'California Vehicle Code ' + cvc_codes + ' "' + cvc_text_code + '".'

    return cvc_text
    

def find_duplicate(lst):
    seen = set()
    for element in lst:
        if element in seen:
            return element.strip()
        seen.add(element)
    return None  # If no duplicates are found

def get_first_names(names):
    first_names = []
    for name in names.split(","):
        first_names.append(name.strip().replace("and ", "").split(" ")[0].strip())
    return first_names

def find_element_index(last_names, duplicate_lastname):
    indexes_found = []
    for i, last_name in enumerate(last_names):
        if last_name == duplicate_lastname:
            indexes_found.append(i)  # Save the index of matching last name
    return indexes_found

def add_names_to_duplicate_lastnames(client_names, mr_mrs_client_last_name):
    last_names = [i.strip() for i in mr_mrs_client_last_name.replace(", and ", ",").split(",")]
    duplicate_lastname = find_duplicate(last_names)

    first_names = get_first_names(client_names)

    for index in find_element_index(last_names, duplicate_lastname):
        last_names[index] = f"{last_names[index]} ({first_names[index]})"

    last_names[-1] = "and " + last_names[-1]
    return ", ".join(last_names)

def custom_title(text, excluded_words=None):
    if excluded_words is None:
        excluded_words = []

    # Create a set of excluded words for O(1) lookup time
    excluded_words_set = set(excluded_words)

    return " ".join(
        word if word.lower() in excluded_words_set else word.capitalize()
        for word in text.split()
    )

def parse_file_data():
    with open(file_template_source, "r", encoding="utf-8") as file:
        return [i.strip().split(":")[1].strip() for i in file.readlines()]
    
def format_currency(amount):

    try:
        # Ensure the input is a float
        amount = float(amount)
        return "${:,.2f}".format(amount)
    except ValueError:
        # Handle the case where the input is not a number
        raise ValueError(f"Invalid input for currency formatting: {amount}")
    
ensure_file_exists(file_template_data, file_template_source)
DATA = parse_file_data()

# Client Information
CLIENT_NAME = DATA[0]
CLIENT_SEX = DATA[1]
IS_YOUNG = DATA[2]
IS_YOUNG = "young" if IS_YOUNG == "yes" else (IS_YOUNG if len(IS_YOUNG) > 1 else "")

# Insured Information
INSURED_NAME = DATA[3]
INSURED_SEX = DATA[4]
INSURED_TITLE = "Mr. " if INSURED_SEX.lower() == "man" else "Mrs. "

# Contact and Claim Information
VIA_TYPE_OPINS = DATA[5]
OPINS = DATA[6]
CLAIM_NUMBER = DATA[7]
DATE_OF_LOSS = DATA[8]
CLAIM_RESPONSIBLE_RECEIVER = DATA[9]

# California Civil Code Text
CALIFORNIA_CVC_TEXT = create_string(DATA[10])

VIA_TYPE_CINS = DATA[11]  # Email CINS
CINS = DATA[12] #CINS NAME 
POLICY_NUMBER = DATA[13] # I.E 23768698269
LIMIT_COVERAGE_CINS = format_currency(DATA[14]) # Format from 3000 -> $3,000.00 


INSURANCE_INIT_OPINS = OPINS.split(" ")[0]
INSURANCE_NAME_CAP_OPINS = OPINS.upper()

INSURANCE_INIT_CINS = CINS.split(" ")[0]
INSURANCE_NAME_CAP_CINS = CINS.upper()

# Automatically set gender-specific variables based on CLIENT_SEX
if CLIENT_SEX == "woman":
    HE_SHE_CLIENT = "she"
    HER_HIM_CLIENT = "her"
    HER_HIS_CLIENT = "her"
    HER_HIS_CLIENT_CAP = "HER"
    CLIENT_TITLE = "Ms. " if IS_YOUNG else "Mrs. "
    HE_SHE_CLIENT_PAGE7 = HE_SHE_CLIENT + " was" 

elif CLIENT_SEX == "man":  # CLIENT_SEX == "man"
    HE_SHE_CLIENT = "he"
    HER_HIM_CLIENT = "him"
    HER_HIS_CLIENT = "his"
    HER_HIS_CLIENT_CAP = "HIS"
    CLIENT_TITLE = "Mr. "
    HE_SHE_CLIENT_PAGE7 = HE_SHE_CLIENT + " was" 

else:
    CLIENT_TITLE = []
    for index, client_sex in enumerate(CLIENT_SEX.split(",")):
        if client_sex.strip() == "woman":
            if IS_YOUNG.split(",")[index].strip() == "no":
                CLIENT_TITLE.append("Mrs.")
            else:
                CLIENT_TITLE.append("Ms.")

        elif client_sex.strip() == "minor":
            CLIENT_TITLE.append("")
        else:
            CLIENT_TITLE.append("Mr.")

    HE_SHE_CLIENT = "they"
    HER_HIM_CLIENT = "them"
    HER_HIS_CLIENT = "their"
    HER_HIS_CLIENT_CAP = "THEIR"

    # Custom
    HE_SHE_CLIENT_PAGE7 = HE_SHE_CLIENT + " were" 
    
# Automatically set gender-specific variables based on INSURED_SEX
if INSURED_SEX == "woman":
    HE_SHE_INSURED = "she"
    HER_HIM_INSURED = "her"
    HER_HIS_INSURED = "her"

elif INSURED_SEX == "man":
    HE_SHE_INSURED = "he"
    HER_HIM_INSURED = "him"
    HER_HIS_INSURED = "his"


# Format the date as MM/DD/YYYY
SETTLEMENT_EXP_DATE = (datetime.now() + relativedelta(months=1)).strftime("%m/%d/%Y")
SETTLEMENT_EXP_DATE = datetime.strptime(SETTLEMENT_EXP_DATE, "%m/%d/%Y").strftime("%B %d, %Y").upper()

SETTLEMENT_EXP_DATE_TITLE = SETTLEMENT_EXP_DATE.title()

CLIENT_NAME_ALL_CAP = CLIENT_NAME.upper()
CLIENT_NAME_EACH_CAP = custom_title(CLIENT_NAME, ["and"])

if " and " not in CLIENT_NAME:
    CLIENT_LAST_NAME = CLIENT_NAME_EACH_CAP.split(" ")[-1] if CLIENT_NAME_EACH_CAP.split(" ")[-1] not in ["Sr", "Jr"] else CLIENT_NAME_EACH_CAP.split(" ")[-2] + " " +CLIENT_NAME_EACH_CAP.split(" ")[-1]
    MR_MRS_CLIENT_NAME_EACH_CAP = (CLIENT_TITLE + CLIENT_NAME_EACH_CAP).title()
    MR_MRS_CLIENT_NAME_ALL_CAP = (CLIENT_TITLE + CLIENT_NAME_EACH_CAP).upper()
    MR_MRS_CLIENT_LAST_NAME = CLIENT_TITLE + CLIENT_LAST_NAME
    
elif " and " in CLIENT_NAME and "," not in CLIENT_NAME: # more than one client 
    MR_MRS_CLIENT_LAST_NAME = ""
    MR_MRS_CLIENT_NAME = ""

    for index, client_name in enumerate(CLIENT_NAME.split(" and ")):
        client_name = client_name.strip()  # Remove extra spaces

        # Add " and " if there's already a name
        if MR_MRS_CLIENT_LAST_NAME:
            MR_MRS_CLIENT_LAST_NAME += ", and "
        if MR_MRS_CLIENT_NAME:
            MR_MRS_CLIENT_NAME += ", and "

        # Check for "Sr" or "Jr" in the name
        if any(title in client_name for title in ["Sr", "Jr"]):
            # Add full name and last two words for titles
            if CLIENT_TITLE[index]:
                MR_MRS_CLIENT_NAME += CLIENT_TITLE[index] + " " + client_name
                MR_MRS_CLIENT_LAST_NAME += (
                    CLIENT_TITLE[index] + " " + " ".join(client_name.split()[-2:])
                )
            else:
                MR_MRS_CLIENT_NAME += client_name
                MR_MRS_CLIENT_LAST_NAME += client_name
        else:
            if CLIENT_TITLE[index]:
                # Add full name and last word (last name)
                MR_MRS_CLIENT_NAME += CLIENT_TITLE[index] + " " + client_name
                MR_MRS_CLIENT_LAST_NAME += CLIENT_TITLE[index] + " " + client_name.split()[-1]
            else:
                MR_MRS_CLIENT_NAME += client_name
                MR_MRS_CLIENT_LAST_NAME += client_name

    MR_MRS_CLIENT_NAME_EACH_CAP = custom_title(MR_MRS_CLIENT_NAME, ["and"])
    MR_MRS_CLIENT_NAME_ALL_CAP = MR_MRS_CLIENT_NAME.upper()

else:
    MR_MRS_CLIENT_LAST_NAME = ""
    MR_MRS_CLIENT_NAME = ""

    for index, client_name in enumerate(CLIENT_NAME.split(", ")):
        last_name_tracker = []
        if "and" in client_name[0:3]:
            client_name = client_name[4:]

        if index != len(CLIENT_NAME.split(", "))-1:
            # Add ", " if there's already a name
            if MR_MRS_CLIENT_LAST_NAME:
                MR_MRS_CLIENT_LAST_NAME += ", "
            if MR_MRS_CLIENT_NAME:
                MR_MRS_CLIENT_NAME += ", "

        else:
            # Add ", " if there's already a name
            if MR_MRS_CLIENT_LAST_NAME:
                MR_MRS_CLIENT_LAST_NAME += ", and "
            if MR_MRS_CLIENT_NAME:
                MR_MRS_CLIENT_NAME += ", and "

    #     # Check for "Sr" or "Jr" in the name
        if any(title in client_name for title in ["Sr", "Jr"]):
            # Add full name and last two words for titles
            if CLIENT_TITLE[index]:
                MR_MRS_CLIENT_NAME += CLIENT_TITLE[index] + " " + client_name
                MR_MRS_CLIENT_LAST_NAME += (
                    CLIENT_TITLE[index] + " " + " ".join(client_name.split()[-2:])
                )
            else:
                MR_MRS_CLIENT_NAME += client_name
                MR_MRS_CLIENT_LAST_NAME += client_name

        else:
            last_name_tracker.append(client_name.split()[-1])
            if CLIENT_TITLE[index]:
                
                MR_MRS_CLIENT_NAME += CLIENT_TITLE[index] + " " + client_name
                MR_MRS_CLIENT_LAST_NAME += CLIENT_TITLE[index] + " " + client_name.split()[-1]

            else:
                MR_MRS_CLIENT_NAME += client_name
                MR_MRS_CLIENT_LAST_NAME += client_name

    MR_MRS_CLIENT_NAME_EACH_CAP = custom_title(MR_MRS_CLIENT_NAME, ["and"])
    MR_MRS_CLIENT_NAME_ALL_CAP = MR_MRS_CLIENT_NAME.upper()


if " and " in MR_MRS_CLIENT_LAST_NAME:
    MR_MRS_CLIENT_LAST_NAME = add_names_to_duplicate_lastnames(CLIENT_NAME, MR_MRS_CLIENT_LAST_NAME)

MR_MRS_CLIENT_LAST_NAME = MR_MRS_CLIENT_LAST_NAME.title()

INSURED_NAME_ALL_CAP = INSURED_NAME.upper()
INSURED_NAME_EACH_CAP = INSURED_NAME.title()

MR_MRS_INSURED_NAME_EACH_CAP = (INSURED_TITLE + INSURED_NAME_ALL_CAP).title()
MR_OR_MRS_INSURED_NAME_ALL_CAP = MR_MRS_INSURED_NAME_EACH_CAP.upper()
DATE_OF_LOSS_FORMATTED = datetime.strptime(DATE_OF_LOSS, "%m/%d/%Y").strftime("%B %d, %Y")

if "," in CLIENT_SEX:
    if "man, man" == CLIENT_SEX:
        CLIENT_SEX = f"healthy men lose"
    elif "woman, woman" == CLIENT_SEX:
        CLIENT_SEX = f"healthy women lose"
    elif "minor" in CLIENT_SEX:
        if "woman" in CLIENT_SEX.split(", ") and "man" in CLIENT_SEX.split(", "):
            CLIENT_SEX = f"healthy men, women, or minors lose"
        elif "woman" in CLIENT_SEX.split(", "):
            CLIENT_SEX = f"healthy women, or minors lose"
        else:
            CLIENT_SEX = f"healthy men, or minors lose"
    else:
        CLIENT_SEX = f"healthy men or women lose"
else:
    CLIENT_SEX = f"a healthy {CLIENT_SEX} loses"

# Store variables in a dictionary
CLIENT_DATA = {
    "CLIENT_NAME": CLIENT_NAME,
    "CLIENT_SEX": CLIENT_SEX,
    "IS_YOUNG": IS_YOUNG,
    "INSURED_NAME": INSURED_NAME,
    "INSURED_SEX": INSURED_SEX,
    "INSURED_TITLE": INSURED_TITLE,
    "HER_HIS_INSURED": HER_HIS_INSURED,
    "HE_SHE_INSURED": HE_SHE_INSURED,
    "VIA_TYPE_OPINS": VIA_TYPE_OPINS,
    "OPINS": OPINS,
    "CLAIM_NUMBER": CLAIM_NUMBER,
    "DATE_OF_LOSS": DATE_OF_LOSS,
    "CLAIM_RESPONSIBLE_RECEIVER": CLAIM_RESPONSIBLE_RECEIVER.title(),
    "CALIFORNIA_CVC_TEXT": CALIFORNIA_CVC_TEXT,
    "INSURANCE_INIT_OPINS": INSURANCE_INIT_OPINS,
    "HE_SHE_CLIENT": HE_SHE_CLIENT,
    "HE_SHE_CLIENT_PAGE7": HE_SHE_CLIENT_PAGE7,
    "HER_HIM_CLIENT": HER_HIM_CLIENT,
    "HER_HIS_CLIENT": HER_HIS_CLIENT,
    "HER_HIS_CLIENT_CAP": HER_HIS_CLIENT_CAP,
    "MR_MRS_CLIENT_NAME_EACH_CAP": MR_MRS_CLIENT_NAME_EACH_CAP,
    "MR_MRS_CLIENT_NAME_ALL_CAP": MR_MRS_CLIENT_NAME_ALL_CAP,
    "SETTLEMENT_EXP_DATE": SETTLEMENT_EXP_DATE,
    "SETTLEMENT_EXP_DATE_TITLE": SETTLEMENT_EXP_DATE_TITLE,
    "CLIENT_NAME_ALL_CAP": CLIENT_NAME_ALL_CAP,
    "CLIENT_NAME_EACH_CAP": CLIENT_NAME_EACH_CAP,
    "MR_MRS_CLIENT_LAST_NAME": MR_MRS_CLIENT_LAST_NAME,
    "INSURED_NAME_ALL_CAP": INSURED_NAME_ALL_CAP,
    "INSURED_NAME_EACH_CAP": INSURED_NAME_EACH_CAP,
    "MR_MRS_INSURED_NAME_EACH_CAP": MR_MRS_INSURED_NAME_EACH_CAP,
    "MR_OR_MRS_INSURED_NAME_ALL_CAP": MR_OR_MRS_INSURED_NAME_ALL_CAP,
    "DATE_OF_LOSS_FORMATTED": DATE_OF_LOSS_FORMATTED,
    "INSURANCE_NAME_CAP_OPINS": INSURANCE_NAME_CAP_OPINS,

    # TEMPLATE UM DATA
    "POLICY_NUMBER": POLICY_NUMBER,
	"LIMIT_COVERAGE_CINS": LIMIT_COVERAGE_CINS,
    "VIA_TYPE_CINS": VIA_TYPE_CINS,
    "CINS": CINS,
    "INSURANCE_INIT_CINS": INSURANCE_INIT_CINS,
    "INSURANCE_NAME_CAP_CINS": INSURANCE_NAME_CAP_CINS,
}


def edit_docx_preserve_format(doc):
    """
    Replaces target text with replacement text in the document while preserving formatting.
    """
    try:
        # Edit document content
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:  # Access individual formatted text segments (runs)
                if CLIENT_DATA.get(run.text) is not None:
                    run.text = run.text.replace(run.text, CLIENT_DATA[run.text])

        # Edit headers
        for section in doc.sections:
            header = section.header  # Access the header for the section
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    if CLIENT_DATA.get(run.text, None):
                        run.text = run.text.replace(run.text, CLIENT_DATA[run.text])

    except Exception as e:
        print(f"An error occurred: {e}")

def draft_document(doc_template, file_type):
    # Paths for the input and output files
    output_path = working_directory / (CLIENT_NAME_ALL_CAP + " - "  + DATE_OF_LOSS_FORMATTED.upper() + f" - ({file_type}).docx")

    if check_and_warn_if_file_exists(output_path):
        # Load the document
        doc = Document(doc_template)

        # Replace placeholders
        edit_docx_preserve_format(doc)

        # Save the updated document
        doc.save(output_path)
        print(f"Document saved as: {output_path}")

if __name__ == "__main__":
    draft_document(template_word_path_opins, "OPINS")
    draft_document(template_word_path_cins_a, "CINS A")
    draft_document(template_word_path_cins_b, "CINS B")